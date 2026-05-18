from __future__ import annotations

import hashlib
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx"}
UNSUPPORTED_EXTENSIONS = {".doc"}

IGNORED_DIRS = {
    ".git",
    ".idea",
    ".venv",
    "__pycache__",
    "data",
    "generated_notes",
    "node_modules",
}


@dataclass(frozen=True)
class SourceFile:
    path: Path
    relative_path: str
    title: str
    size: int
    extension: str


@dataclass(frozen=True)
class Topic:
    id: str
    title: str
    content: str
    level: int = 0


@dataclass(frozen=True)
class PdfPage:
    number: int
    text: str


@dataclass(frozen=True)
class DocumentContent:
    title: str
    content: str
    pages: list[PdfPage] | None = None


def list_source_files(root: Path, recursive: bool = False) -> list[SourceFile]:
    if not root.exists() or not root.is_dir():
        return []

    pattern = "**/*" if recursive else "*"
    files: list[SourceFile] = []
    for path in sorted(root.glob(pattern)):
        if not path.is_file() or _is_ignored(path, root):
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        files.append(_build_source_file(path, root))
    return files


def read_document(path: Path) -> DocumentContent:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        text = read_text_document(path)
        return DocumentContent(title=extract_title(text, path.stem), content=text)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        text = read_docx(path)
        return DocumentContent(title=extract_title(text, path.stem), content=text)
    if suffix in UNSUPPORTED_EXTENSIONS:
        raise ValueError("暂不支持 .doc 老 Word 格式，请先另存为 .docx。")
    raise ValueError(f"暂不支持这种文件格式：{suffix or '无扩展名'}")


def read_uploaded_document(file_name: str, data: bytes) -> DocumentContent:
    suffix = Path(file_name).suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        text = decode_text(data)
        return DocumentContent(title=extract_title(text, Path(file_name).stem), content=text)
    if suffix == ".pdf":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            return read_pdf(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)
    if suffix == ".docx":
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            text = read_docx(tmp_path)
            return DocumentContent(title=extract_title(text, Path(file_name).stem), content=text)
        finally:
            tmp_path.unlink(missing_ok=True)
    if suffix in UNSUPPORTED_EXTENSIONS:
        raise ValueError("暂不支持 .doc 老 Word 格式，请先另存为 .docx。")
    raise ValueError(f"暂不支持这种文件格式：{suffix or '无扩展名'}")


def read_text_document(path: Path) -> str:
    return decode_text(path.read_bytes())


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def read_pdf(path: Path) -> DocumentContent:
    reader = PdfReader(str(path))
    pages: list[PdfPage] = []
    page_texts: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        pages.append(PdfPage(number=index, text=text))
        if text:
            page_texts.append(f"[第 {index} 页]\n\n{text}")

    combined = "\n\n".join(page_texts).strip()
    return DocumentContent(title=path.stem, content=combined, pages=pages)


def read_docx(path: Path) -> str:
    document = DocxDocument(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts).strip()


def extract_title(content: str, fallback: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip() or fallback
        return stripped[:80]
    return fallback


def split_topics(content: str) -> list[Topic]:
    if not content.strip():
        return [Topic(id="whole", title="整篇文档", content="", level=0)]

    topics = [Topic(id="whole", title="整篇文档", content=content, level=0)]
    headings = list(re.finditer(r"^(#{2,3})\s+(.+?)\s*$", content, flags=re.MULTILINE))

    for index, match in enumerate(headings):
        start = match.start()
        end = headings[index + 1].start() if index + 1 < len(headings) else len(content)
        level = len(match.group(1))
        title = match.group(2).strip()
        body = content[start:end].strip()
        topic_id = f"h{level}-{index}-{_slug(title)}"
        topics.append(Topic(id=topic_id, title=title, content=body, level=level))

    return topics


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()[:16]


def _build_source_file(path: Path, root: Path) -> SourceFile:
    try:
        if path.suffix.lower() in {".md", ".markdown", ".txt"}:
            text = read_text_document(path)
            title = extract_title(text, path.stem)
        else:
            title = path.stem
    except Exception:
        title = path.stem
    return SourceFile(
        path=path,
        relative_path=str(path.relative_to(root)),
        title=title,
        size=path.stat().st_size,
        extension=path.suffix.lower(),
    )


def _slug(text: str) -> str:
    slug = re.sub(r"\s+", "-", text.strip().lower())
    slug = re.sub(r"[^\w\-\u4e00-\u9fff]+", "", slug)
    return slug[:48] or "topic"


def _is_ignored(path: Path, root: Path) -> bool:
    try:
        parts = path.relative_to(root).parts
    except ValueError:
        return False
    return any(part in IGNORED_DIRS for part in parts)
